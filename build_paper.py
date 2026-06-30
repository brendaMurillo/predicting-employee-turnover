from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup: US Letter, narrow margins ────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(0.75)
section.right_margin  = Inches(0.75)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Two-column layout via sectPr XML ─────────────────────────────────────────
def set_two_columns(section):
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')  # 0.5 inch gutter
    sectPr.append(cols)

set_two_columns(section)

# ── Helper: add a run with formatting ────────────────────────────────────────
def add_run(para, text, bold=False, italic=False, size=None, font_name='Times New Roman'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    run.font.name = font_name
    return run

def set_para_font(para, size=10, font_name='Times New Roman'):
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.name = font_name

# ── TITLE (centered, large, serif) ───────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title_para,
        'Predicting Employee Turnover: A Statistical Analysis\nof Workforce Attrition Factors',
        bold=True, size=18)
title_para.paragraph_format.space_after = Pt(12)

# ── AUTHOR BLOCK ─────────────────────────────────────────────────────────────
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(author, 'Brenda Murillo\n', bold=False, size=11)
add_run(author, 'Dept. of Computer and Electrical Engineering and Computer Science\n', size=10, italic=True)
add_run(author, 'California State University, Bakersfield\n', size=10, italic=True)
add_run(author, 'Bakersfield, CA, USA\n', size=10)
add_run(author, 'bmurillo3@csub.edu', size=10)
author.paragraph_format.space_after = Pt(10)

# ── Section heading helper (centered, bold, Roman numeral style) ──────────────
def section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f'{number}.  {title.upper()}', bold=True, size=10)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    return p

# ── Subsection heading helper (left, italic letter) ───────────────────────────
def sub_heading(doc, letter, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, f'{letter}.  ', bold=False, italic=True, size=10)
    add_run(p, title, bold=False, italic=True, size=10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    return p

# ── Body paragraph helper ─────────────────────────────────────────────────────
def body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.2)
    add_run(p, text, size=10)
    p.paragraph_format.space_after = Pt(4)
    return p

# ── ABSTRACT ─────────────────────────────────────────────────────────────────
abs_para = doc.add_paragraph()
abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(abs_para, 'Abstract—', bold=True, italic=True, size=10)
add_run(abs_para,
    'Employee turnover represents one of the most persistent and costly '
    'challenges facing organizations today, yet most retention strategies rely '
    'on intuition rather than data. This paper presents a statistical analysis '
    'of workforce attrition using an HR dataset of 14,999 employee records drawn '
    'from Kaggle. We apply descriptive statistics, relational database querying '
    '(SQLite), correlation analysis, and visual analysis to identify the factors '
    'most strongly associated with employee departure. Our findings reveal an '
    'overall turnover rate of 23.81%, with low-salary employees leaving at nearly '
    '30%—more than four times the rate of high-salary employees. Job '
    'satisfaction emerges as the single strongest predictor (r = −0.388); '
    'employees who left averaged a satisfaction score of 0.44 compared to 0.67 '
    'for those who stayed. We further identify 615 current employees who match a '
    'high-risk churn profile. These findings provide HR departments with '
    'actionable, data-driven criteria for early intervention and targeted '
    'retention strategies.',
    bold=True, italic=False, size=10)
abs_para.paragraph_format.space_after = Pt(4)

# Index Terms
idx = doc.add_paragraph()
idx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(idx, 'Index Terms—', bold=True, italic=True, size=10)
add_run(idx,
    'employee turnover, workforce attrition, HR analytics, churn prediction, '
    'descriptive statistics, SQLite, data analysis',
    bold=True, size=10)
idx.paragraph_format.space_after = Pt(8)

# ─────────────────────────────────────────────────────────────────────────────
# I. INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, 'I', 'Introduction')
body(doc,
    'The cost of replacing an employee is widely estimated at between 50% and '
    '200% of that employee\'s annual salary, accounting for recruiting, '
    'onboarding, lost productivity, and institutional knowledge transfer [2]. '
    'Yet despite this substantial burden, many organizations respond to turnover '
    'reactively—conducting exit interviews only after employees have already '
    'resigned, and implementing changes too late to retain the individuals who '
    'prompted them.')
body(doc,
    'The central challenge is one of prediction: given observable workplace '
    'factors, can an organization identify which employees are most likely to '
    'leave before they submit their resignation? This problem is well-suited to '
    'data-driven analysis. Modern HR systems routinely capture variables such as '
    'satisfaction surveys, performance evaluations, compensation tiers, workload '
    'metrics, and promotion histories—precisely the kinds of signals that '
    'may indicate attrition risk.')
body(doc,
    'This paper addresses that challenge by applying statistical analysis to an '
    'HR dataset containing records for 14,999 employees. We pursue four core '
    'research questions: (1) How strongly does salary level predict turnover? '
    '(2) What is the relationship between job satisfaction and departure? '
    '(3) Does workload influence turnover? '
    '(4) Can we identify current employees at high churn risk using observable '
    'characteristics?')
body(doc,
    'Our approach is grounded in accessible statistical methods—descriptive '
    'statistics, group comparisons, correlation analysis, and SQL querying against '
    'a normalized relational database—rather than black-box machine learning '
    'models, making our findings interpretable and actionable for HR practitioners.')

# ─────────────────────────────────────────────────────────────────────────────
# II. RELATED WORK
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, 'II', 'Related Work')
body(doc,
    'Employee attrition prediction has been an active area of research in both '
    'human resources management and data science. Early work by Mobley [1] '
    'established a foundational behavioral model of turnover, identifying job '
    'dissatisfaction as the primary precursor to resignation. Griffeth, Hom, and '
    'Gaertner [2] confirmed through meta-analysis that satisfaction, pay, and '
    'workload are the most reliable predictors of departure across industries.')
body(doc,
    'More recent computational approaches have applied supervised machine learning '
    'to HR datasets. Punnoose and Ajit [5] applied random forests and logistic '
    'regression to the same Kaggle HR dataset, achieving accuracies above 97%. '
    'Alduayj and Rajpoot [7] compared decision trees, SVMs, and neural networks, '
    'finding tree-based models consistently outperformed others on HR tabular data. '
    'Zhao et al. [6] demonstrated ensemble methods improve performance but at '
    'significant cost to interpretability [8].')
body(doc,
    'Our work is positioned in the tradition of transparent statistical approaches '
    'that make decision criteria visible and reproducible by any analyst, '
    'prioritizing interpretability over predictive accuracy.')

# ─────────────────────────────────────────────────────────────────────────────
# III. DATASET
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, 'III', 'Dataset')

sub_heading(doc, 'A', 'Source and Scope')
body(doc,
    'The dataset is the HR Employee Retention dataset publicly available on '
    'Kaggle [9], compiled to simulate realistic human resources records. It '
    'contains 14,999 employee records, each representing a unique individual. '
    'The outcome variable (left) indicates whether the employee departed (1) '
    'or remained (0).')

sub_heading(doc, 'B', 'Variables')
body(doc, 'The dataset includes ten variables spanning four conceptual domains, shown in Table I.', indent=True)

# TABLE I
t1_caption = doc.add_paragraph('TABLE I.   DATASET VARIABLES')
t1_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t1_caption.runs:
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

t1 = doc.add_table(rows=10, cols=4)
t1.style = 'Table Grid'
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Domain', 'Variable', 'Type', 'Description']
rows_data = [
    ['Satisfaction',  'satisfaction_level',    'Numeric [0,1]', 'Self-reported satisfaction score'],
    ['Performance',   'last_evaluation',        'Numeric [0,1]', 'Most recent performance review'],
    ['Workload',      'number_project',          'Integer',       'Number of projects assigned'],
    ['Workload',      'average_monthly_hours',   'Integer',       'Avg. hours worked per month'],
    ['Tenure',        'time_spend_company',      'Integer',       'Years employed at company'],
    ['Safety',        'work_accident',           'Boolean',       'Had a work accident'],
    ['Career',        'promotion_last_5years',   'Boolean',       'Promoted in last 5 years'],
    ['Compensation',  'salary',                  'Categorical',   'Low / medium / high'],
    ['Outcome',       'left',                    'Boolean',       'Left the company'],
]
for i, h in enumerate(headers):
    cell = t1.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
for r_idx, row_data in enumerate(rows_data):
    for c_idx, val in enumerate(row_data):
        cell = t1.rows[r_idx + 1].cells[c_idx]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
doc.add_paragraph()

sub_heading(doc, 'C', 'Data Quality')
body(doc,
    'A missing value audit revealed no null entries in any column. Outlier '
    'detection using the IQR method identified anomalous values in '
    'average_monthly_hours (above 310 hrs/month), retained as plausible '
    'extreme-workload cases. Schema constraints enforced CHECK ranges on '
    'satisfaction and evaluation scores ([0,1]), non-negative hours and tenure, '
    'and foreign key consistency across all four tables. A row-count '
    'reconciliation confirmed all 14,999 records are present after normalization.')

sub_heading(doc, 'D', 'Database Schema')
body(doc,
    'The flat CSV was normalized into a four-table relational schema in SQLite [10]: '
    '(1) salary_levels—lookup table mapping salary tier labels to IDs; '
    '(2) employees—core identity table (emp_id, salary_id, tenure); '
    '(3) performance_records—satisfaction, evaluation, project count, monthly hours; '
    '(4) employment_events—work accidents, promotion history, and turnover outcome. '
    'This design eliminates redundancy, enforces referential integrity via foreign '
    'keys, and allows the turnover outcome to be joined against any combination '
    'of attributes in a single SQL query.')

# ─────────────────────────────────────────────────────────────────────────────
# IV. METHODOLOGY
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, 'IV', 'Methodology')

sub_heading(doc, 'A', 'Descriptive Statistics')
body(doc,
    'For all numeric variables, we computed mean, median, mode, standard '
    'deviation, variance, range, Q1, Q3, and the interquartile range, '
    'calculated separately for employees who left and those who stayed.')

sub_heading(doc, 'B', 'Group Comparison')
body(doc,
    'Turnover rates were computed for each level of categorical variables '
    '(salary tier, promotion status, work accident). For numeric variables, '
    'employees were segmented into bins and turnover rates computed within each '
    'bin to identify non-linear relationships.')

sub_heading(doc, 'C', 'Correlation Analysis')
body(doc,
    'Pearson correlation coefficients were computed between each numeric feature '
    'and the binary left outcome, producing a ranked list of predictors by '
    'association strength.')

sub_heading(doc, 'D', 'SQL-Based Analysis')
body(doc,
    'All group-level queries were executed against the normalized SQLite database, '
    'including overall turnover rate, turnover by salary tier (three-table join), '
    'average satisfaction and monthly hours by turnover status, and high-risk '
    'employee identification using compound WHERE filters.')

sub_heading(doc, 'E', 'Risk Profiling')
body(doc,
    'A high-risk profile was defined using three simultaneous criteria: '
    'satisfaction score below 0.4, salary tier of "low", and no promotion in '
    'the last five years, constrained to employees still at the company. '
    'This threshold was set based on the distributional boundary at which '
    'turnover rates were observed to increase sharply.')

# ─────────────────────────────────────────────────────────────────────────────
# V. RESULTS
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, 'V', 'Results')

sub_heading(doc, 'A', 'Overall Turnover Rate')
body(doc,
    'Of 14,999 employees, 3,571 left the organization, yielding an overall '
    'turnover rate of 23.81%. This substantially exceeds the commonly cited '
    'U.S. industry average of approximately 15% [4].')

sub_heading(doc, 'B', 'Turnover by Salary Tier')
body(doc,
    'Salary tier is one of the strongest single-variable predictors of turnover. '
    'Table II shows that low-salary employees leave at more than four times the '
    'rate of high-salary employees.')

# TABLE II
t2_caption = doc.add_paragraph('TABLE II.   TURNOVER RATE BY SALARY TIER')
t2_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t2_caption.runs:
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

t2 = doc.add_table(rows=4, cols=4)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
t2_data = [
    ['Salary Tier', 'Total Employees', 'Left', 'Turnover Rate'],
    ['Low',    '7,316', '2,172', '29.69%'],
    ['Medium', '6,446', '1,317', '20.43%'],
    ['High',   '1,237',    '82',  '6.63%'],
]
for r_idx, row_data in enumerate(t2_data):
    for c_idx, val in enumerate(row_data):
        cell = t2.rows[r_idx].cells[c_idx]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.bold = (r_idx == 0)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
doc.add_paragraph()

sub_heading(doc, 'C', 'Job Satisfaction and Turnover')
body(doc,
    'Job satisfaction shows the most pronounced difference between retained and '
    'departed employees (Table III). It produces the strongest Pearson correlation '
    'with turnover (r = −0.388). Satisfaction differs by 0.227 '
    'points between groups—a 34% relative gap. By contrast, performance '
    'evaluation scores are nearly identical (0.715 vs. 0.718), confirming that '
    'turnover is a satisfaction problem, not a performance problem.')

# TABLE III
t3_caption = doc.add_paragraph('TABLE III.   AVERAGE FEATURE VALUES BY TURNOVER STATUS')
t3_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t3_caption.runs:
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

t3 = doc.add_table(rows=6, cols=3)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
t3_data = [
    ['Feature', 'Stayed', 'Left'],
    ['Satisfaction level',    '0.667', '0.440'],
    ['Last evaluation score', '0.715', '0.718'],
    ['Number of projects',    '3.79',  '3.86'],
    ['Avg. monthly hours',    '199.1', '207.4'],
    ['Years at company',      '3.38',  '3.88'],
]
for r_idx, row_data in enumerate(t3_data):
    for c_idx, val in enumerate(row_data):
        cell = t3.rows[r_idx].cells[c_idx]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.bold = (r_idx == 0)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
doc.add_paragraph()

sub_heading(doc, 'D', 'Feature Correlation Analysis')
body(doc,
    'Table IV presents the Pearson correlation of each feature with the left '
    'outcome, ranked by absolute magnitude. Satisfaction is the strongest '
    'predictor by a wide margin (r = −0.388). last_evaluation '
    'shows virtually no correlation (r = +0.007), confirming that '
    'performance evaluation does not distinguish leavers from stayers.')

# TABLE IV
t4_caption = doc.add_paragraph('TABLE IV.   FEATURE CORRELATION WITH TURNOVER')
t4_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t4_caption.runs:
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

t4 = doc.add_table(rows=8, cols=2)
t4.style = 'Table Grid'
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
t4_data = [
    ['Feature', 'Pearson r'],
    ['satisfaction_level',     '−0.388'],
    ['work_accident',          '−0.155'],
    ['time_spend_company',     '+0.145'],
    ['average_monthly_hours',  '+0.071'],
    ['promotion_last_5years',  '−0.062'],
    ['number_project',         '+0.024'],
    ['last_evaluation',        '+0.007'],
]
for r_idx, row_data in enumerate(t4_data):
    for c_idx, val in enumerate(row_data):
        cell = t4.rows[r_idx].cells[c_idx]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.bold = (r_idx == 0)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
doc.add_paragraph()

sub_heading(doc, 'E', 'Workload and Turnover')
body(doc,
    'Employees who left worked an average of 207.4 hours per month compared to '
    '199.1 for those who stayed—a difference of 8.3 hours, or roughly two '
    'extra workdays per month. The correlation of average_monthly_hours with '
    'turnover (r = +0.071) suggests overwork is a contributing factor, '
    'though it acts as a compounding stressor alongside low satisfaction and low '
    'pay rather than an independent primary driver.')

sub_heading(doc, 'F', 'High-Risk Employee Profiling')
body(doc,
    'Applying the compound risk profile (satisfaction < 0.4, salary = low, no '
    'promotion in 5 years, still employed), the analysis identified 615 current '
    'employees who match all three criteria. Table V shows the ten highest-risk '
    'individuals by lowest satisfaction score.')

# TABLE V
t5_caption = doc.add_paragraph('TABLE V.   SAMPLE HIGH-RISK EMPLOYEES (TOP 10)')
t5_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t5_caption.runs:
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

t5 = doc.add_table(rows=11, cols=4)
t5.style = 'Table Grid'
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
t5_data = [
    ['Emp. ID', 'Satisfaction', 'Hrs/Mo', 'Tenure (yrs)'],
    ['2087',  '0.12', '244', '5'],
    ['3129',  '0.12', '257', '6'],
    ['4026',  '0.12', '241', '2'],
    ['4684',  '0.12', '276', '4'],
    ['5336',  '0.12', '166', '3'],
    ['7772',  '0.12', '161', '4'],
    ['8745',  '0.12', '287', '4'],
    ['9283',  '0.12', '200', '3'],
    ['11069', '0.12', '191', '5'],
    ['13280', '0.12', '191', '5'],
]
for r_idx, row_data in enumerate(t5_data):
    for c_idx, val in enumerate(row_data):
        cell = t5.rows[r_idx].cells[c_idx]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.bold = (r_idx == 0)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# VI. DISCUSSION
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, 'VI', 'Discussion')

sub_heading(doc, 'A', 'Compensation is the Dominant Structural Factor')
body(doc,
    'The four-to-one ratio in turnover rates between low and high salary tiers '
    'is the single most actionable finding in this study. It suggests that '
    'compensation structure creates a systemic attrition risk that cannot be '
    'resolved through engagement programs or managerial coaching alone. '
    'Organizations with large proportions of low-salary employees should treat '
    'compensation review as a primary retention strategy [2].')

sub_heading(doc, 'B', 'Satisfaction Predicts Turnover Independently of Performance')
body(doc,
    'The near-identical evaluation scores between leavers and stayers '
    '(0.718 vs. 0.715), combined with the near-zero correlation of '
    'last_evaluation (r = +0.007), challenges the common assumption '
    'that turnover is driven by underperformers. Departed employees were equally '
    'well-evaluated as retained ones. This replicates Mobley\'s foundational '
    'finding [1] and underscores that retention efforts should be framed as a '
    'satisfaction and engagement problem, not a performance management problem.')

sub_heading(doc, 'C', 'Workload Acts as a Compounding Risk Factor')
body(doc,
    'The highest-risk employees combine low satisfaction with above-average hours, '
    'suggesting overwork amplifies the effect of low satisfaction rather than '
    'driving attrition independently. HR interventions should target the '
    'intersection of these factors, not either factor in isolation [3].')

sub_heading(doc, 'D', 'Early Intervention Opportunity')
body(doc,
    'The identification of 615 high-risk current employees is the most '
    'practically valuable output of this analysis. Unlike post-departure '
    'analyses, this profiling approach is actionable before attrition occurs. '
    'HR departments can operationalize it as a recurring SQL query to generate '
    'a prioritized list for manager outreach, salary review, or promotion '
    'consideration.')

sub_heading(doc, 'E', 'Limitations')
body(doc,
    'This study has several limitations. First, the dataset is simulated [9] '
    'and does not represent a specific organization, limiting generalizability. '
    'Second, the analysis is descriptive rather than causal: we identify '
    'associations, not causal effects. Third, the dataset lacks department or '
    'role information, which likely moderates many relationships observed. '
    'Fourth, class imbalance (~24% leavers) should be addressed before '
    'extending this framework to predictive modeling.')

# ─────────────────────────────────────────────────────────────────────────────
# VII. CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────
section_heading(doc, 'VII', 'Conclusion')
body(doc,
    'This paper analyzed an HR dataset of 14,999 employees to identify the '
    'factors most strongly associated with workforce attrition. Salary tier and '
    'job satisfaction are the two dominant predictors of turnover—with '
    'satisfaction producing the strongest correlation (r = −0.388)'
    '—while performance evaluation is essentially uncorrelated with departure '
    '(r = +0.007). Critically, departed employees were equally '
    'well-evaluated as retained employees, confirming that turnover is a '
    'satisfaction problem, not a performance problem.')
body(doc,
    'By normalizing a flat HR CSV into a constraint-validated SQLite schema and '
    'operationalizing findings as reusable SQL queries, we demonstrate that '
    'organizations can generate actionable, prioritized intervention lists from '
    'existing HR data without requiring machine learning infrastructure. '
    'The 615 high-risk employees identified here represent an immediately '
    'addressable retention opportunity.')
body(doc,
    'Future work should apply this framework to real organizational data with '
    'department-level granularity, incorporate time-series analysis to model '
    'attrition risk over tenure, and explore causal inference methods to move '
    'beyond association toward actionable causal claims about the effect of '
    'salary increases or promotion on retention probability.')

# ── ACKNOWLEDGMENT ────────────────────────────────────────────────────────────
ack_heading = doc.add_paragraph()
ack_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(ack_heading, 'ACKNOWLEDGMENT', bold=True, size=10)
ack_heading.paragraph_format.space_before = Pt(8)
ack_heading.paragraph_format.space_after  = Pt(4)

body(doc,
    'The author thanks the faculty of the Department of Computer and Electrical '
    'Engineering and Computer Science at California State University, Bakersfield '
    'for guidance on this project.')

# ── REFERENCES ────────────────────────────────────────────────────────────────
ref_heading = doc.add_paragraph()
ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(ref_heading, 'REFERENCES', bold=True, size=10)
ref_heading.paragraph_format.space_before = Pt(8)
ref_heading.paragraph_format.space_after  = Pt(4)

refs = [
    '[1] W. H. Mobley, "Intermediate linkages in the relationship between job satisfaction and employee turnover," J. Appl. Psychol., vol. 62, no. 2, pp. 237–240, 1977.',
    '[2] R. W. Griffeth, P. W. Hom, and S. Gaertner, "A meta-analysis of antecedents and correlates of employee turnover," J. Manag., vol. 26, no. 3, pp. 463–488, 2000.',
    '[3] P. W. Hom, T. W. Lee, J. D. Shaw, and J. P. Hausknecht, "One hundred years of employee turnover theory and research," J. Appl. Psychol., vol. 102, no. 3, pp. 530–545, 2017.',
    '[4] Bureau of Labor Statistics, "Job Openings and Labor Turnover Summary," U.S. Department of Labor, 2023.',
    '[5] R. Punnoose and P. Ajit, "Prediction of employee turnover in organizations using machine learning algorithms," Int. J. Adv. Res. Artif. Intell., vol. 5, no. 9, pp. 22–26, 2016.',
    '[6] Y. Zhao, M. K. Hryniewicki, F. Cheng, B. Fu, and X. Zhu, "Employee turnover prediction with machine learning: A reliable approach," in Proc. IntelliSys, 2018, pp. 737–758.',
    '[7] S. S. Alduayj and K. Rajpoot, "Predicting employee attrition using machine learning," in Proc. IIT, 2018, pp. 93–98.',
    '[8] F. Doshi-Velez and B. Kim, "Towards a rigorous science of interpretable machine learning," arXiv:1702.08608, 2017.',
    '[9] HR Analytics / Employee Retention Dataset, Kaggle. [Online]. Available: https://www.kaggle.com/',
    '[10] E. F. Codd, "A relational model of data for large shared data banks," Commun. ACM, vol. 13, no. 6, pp. 377–387, 1970.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, ref, size=9)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/brendamurillo/Desktop/data analyst project/final_paper.docx'
doc.save(out)
print(f'Saved: {out}')
